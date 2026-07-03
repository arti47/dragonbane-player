#!/usr/bin/env python3
"""
Generate the Dragonbane Player App — Beginner's Manual as a .docx.

Run:  python3 docs/build_manual.py
Output: docs/Dragonbane-Player-Beginners-Manual.docx

The content is written specifically for players who are new to BOTH the app and
the Dragonbane RPG. It walks the actual on-screen flow of the app (the five nav
tabs, the 10-step creation wizard, and every panel of the character sheet), and
mirrors the terminology, button labels, and mechanics found in the source
(src/*.js, data.js). Nothing here is invented — labels match the running app.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Theme colours (Dragonbane parchment/ink) -----------------------------
INK      = RGBColor(0x2b, 0x22, 0x1b)   # near-black brown
ACCENT   = RGBColor(0x7a, 0x2e, 0x1d)   # dragonbane red-brown
ACCENT2  = RGBColor(0x9c, 0x6b, 0x2f)   # gold-brown
MUTED    = RGBColor(0x6b, 0x5d, 0x4f)
GOOD     = RGBColor(0x2e, 0x6b, 0x2e)
BAD      = RGBColor(0x9c, 0x2e, 0x2e)

doc = Document()

# ---- Base styles ----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for lvl, (size, color, before) in {
    "Heading 1": (18, ACCENT,  16),
    "Heading 2": (14, ACCENT,  12),
    "Heading 3": (12, ACCENT2, 8),
    "Heading 4": (11, INK,     6),
}.items():
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.keep_with_next = True


# ---- Helpers --------------------------------------------------------------
def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def para(text="", *, italic=False, bold=False, size=None, color=None,
         align=None, space_after=None, space_before=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        r.italic = italic
        r.bold = bold
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return p


def rich(parts, *, style=None, align=None, space_after=None):
    """parts: list of (text, {kwargs}) tuples for mixed formatting in one line."""
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    for text, kw in parts:
        r = p.add_run(text)
        r.bold = kw.get("bold", False)
        r.italic = kw.get("italic", False)
        if kw.get("size"):
            r.font.size = Pt(kw["size"])
        if kw.get("color"):
            r.font.color.rgb = kw["color"]
        if kw.get("mono"):
            r.font.name = "Consolas"
    return p


def bullet(text, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p


def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p


def callout(title, body, tint="F1E9DA", edge=ACCENT):
    """A single-cell shaded box used for tips / notes / warnings."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    _shade(cell, tint)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = edge
    if body:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.add_run(body)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def table(headers, rows, widths=None, header_fill="7A2E1D"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i], header_fill)
        pp = hdr[i].paragraphs[0]
        pp.paragraph_format.space_after = Pt(1)
        run = pp.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1:
                _shade(cells[i], "F5EFE2")
            pp = cells[i].paragraphs[0]
            pp.paragraph_format.space_after = Pt(1)
            run = pp.add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def hrule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "9C6B2F")
    pbdr.append(bottom)
    pPr.append(pbdr)


def pagebreak():
    doc.add_page_break()


# ==========================================================================
# COVER
# ==========================================================================
for _ in range(4):
    doc.add_paragraph()
para("⚔", size=54, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("Dragonbane Player", size=34, bold=True, color=ACCENT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("The Complete Beginner's Manual", size=18, color=ACCENT2, italic=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
para("A step-by-step guide to creating and playing a hero with the "
     "Dragonbane Player app — written for players who are new to the app "
     "and to the Dragonbane roleplaying game.",
     size=12, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(6):
    doc.add_paragraph()
para("For phone · browser · computer  •  Works offline  •  No account required",
     size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Unofficial personal-use companion tool. Dragonbane is © Free League "
     "Publishing.", size=8, color=MUTED, italic=True,
     align=WD_ALIGN_PARAGRAPH.CENTER)
pagebreak()

# ==========================================================================
# TABLE OF CONTENTS (manual)
# ==========================================================================
doc.add_heading("What's in this manual", level=1)
toc = [
    ("1.  Welcome — what this app is", "The app at a glance, and what you need."),
    ("2.  Getting started", "Opening the app, installing it, and the screen layout."),
    ("3.  Dragonbane in 5 minutes", "The core rules you need before you play."),
    ("4.  Creating your first hero", "The 10-step creation wizard, one step at a time."),
    ("5.  Using a ready-made hero", "The five Core Set pre-generated characters."),
    ("6.  Reading your character sheet", "Every panel of the sheet explained."),
    ("7.  Rolling dice", "Skill rolls, Dragons & Demons, boons/banes, pushing."),
    ("8.  Fighting: the Combat tab", "Running a fight with initiative and monsters."),
    ("9.  Rest, healing & getting better", "Recovering and advancing your skills."),
    ("10. Magic", "Learning spells and casting them."),
    ("11. Solo play", "Playing without a Game Master."),
    ("12. The Rules Library", "Looking things up inside the app."),
    ("13. Settings, saving & playing together", "Themes, backups, and cloud sync."),
    ("14. Quick reference", "Tables, a glossary, and beginner tips."),
]
for name, desc in toc:
    rich([(name + "  ", {"bold": True}), (desc, {"italic": True, "color": MUTED, "size": 10})],
         space_after=3)
pagebreak()

# ==========================================================================
# 1. WELCOME
# ==========================================================================
doc.add_heading("1.  Welcome — what this app is", level=1)
para("The Dragonbane Player app is a digital character sheet and dice roller "
     "for the Dragonbane fantasy roleplaying game. It does the fiddly parts "
     "for you: it builds a legal character step by step, tracks your hit "
     "points and gear, rolls the dice with the right bonuses already applied, "
     "and remembers everything between sessions.")

para("You do NOT need to be a rules expert to use it.", bold=True)
para("The app knows the Dragonbane core rules and applies them automatically. "
     "When you tap a skill, it already knows your chance to succeed. When you "
     "swing a weapon, it already knows your damage bonus. Your job is to make "
     "choices and tell the story; the app handles the maths.")

doc.add_heading("What you get", level=3)
bullet("A guided character-creation wizard that can't produce an illegal hero.", bold_lead="")
bullet("A living character sheet — hit points, willpower, conditions, skills, spells, inventory and money, all interactive.", bold_lead="")
bullet("A dice engine that handles Dragons (criticals), Demons (fumbles), boons, banes and 'pushing' a roll.", bold_lead="")
bullet("A combat tracker with a built-in bestiary of monsters and NPCs.", bold_lead="")
bullet("A solo-play assistant, a searchable rules library, and optional cloud sync to share a party across devices.", bold_lead="")

doc.add_heading("What you need to play", level=3)
bullet("This app (on a phone, tablet, or computer).", bold_lead="")
bullet("The Dragonbane rulebook is helpful but not required — the app has a Rules Library built in.", bold_lead="")
bullet("Friends and a Game Master (GM) for a normal game — or nobody at all, if you use Solo Mode.", bold_lead="")

callout("No sign-up needed.",
        "The app runs the moment you open it. Your heroes are saved on your own "
        "device automatically. You only create an account if you later want to "
        "sync a shared party across several devices (see Section 13).",
        tint="EAF2E6", edge=GOOD)

# ==========================================================================
# 2. GETTING STARTED
# ==========================================================================
doc.add_heading("2.  Getting started", level=1)

doc.add_heading("Opening the app", level=2)
para("Open the app's web address in your browser (or the file your group was "
     "given). It loads instantly and works even with no internet connection "
     "after the first visit. In the top-left you'll see the Dragonbane logo; "
     "in the top-right a small pill that reads ", space_after=2)
rich([("Local", {"mono": True, "bold": True}),
      ("  — this tells you your heroes are stored privately on this device — "
       "and a ", {}),
      ("☾ / ☀", {"mono": True, "bold": True}),
      (" button that switches between light and dark themes.", {})])

doc.add_heading("Installing it (optional but recommended)", level=2)
para("Because it is a Progressive Web App (PWA), you can install it like a "
     "normal app so it opens full-screen and works offline:")
bullet("On a phone: open the browser menu and choose 'Add to Home Screen' (or 'Install app').", bold_lead="")
bullet("On a computer (Chrome/Edge): click the install icon in the address bar, or use the browser menu's 'Install' option.", bold_lead="")
para("Once installed it behaves like any other app on your device. Updates "
     "download automatically; if a new version is ready you'll see a small "
     "'Update available — reload' banner.", italic=True, color=MUTED)

doc.add_heading("The screen layout", level=2)
para("Everything is reached from the row of tabs along the bottom of the "
     "screen. There are five (sometimes six) of them:")
table(
    ["Tab", "Icon", "What it's for"],
    [
        ["Heroes", "⚔", "Your list of characters. Create, open, or delete heroes here. This is your home base."],
        ["Combat", "🛡", "The combat tracker: initiative order, monsters, and attack rolls for a fight."],
        ["Solo", "🧭", "The solo-play assistant (oracle, inspiration tables, foe generator). Hidden unless Solo Mode is on."],
        ["Rules", "📖", "A searchable library of all the game's rules, kin, professions, spells and gear."],
        ["About", "⚙", "Settings: themes, content toggles, backup/restore, and cloud-sync sign-in."],
        ["GM", "🎲", "An optional Game Master dashboard. Hidden unless you turn it on or you run a synced campaign."],
    ],
    widths=[1.0, 0.6, 4.6],
)
callout("Tip — start on the Heroes tab.",
        "Whenever you're unsure where you are, tap Heroes (⚔) at the bottom "
        "left to get back to the main screen.")

# ==========================================================================
# 3. DRAGONBANE IN 5 MINUTES
# ==========================================================================
doc.add_heading("3.  Dragonbane in 5 minutes", level=1)
para("Here are the only rules you need to understand before you start. The app "
     "does all of this for you — this section just explains what it's doing.")

doc.add_heading("Your six attributes", level=2)
para("Every hero has six attributes, each scored from 3 to 18. Higher is "
     "better. They are:")
table(
    ["Code", "Attribute", "Governs"],
    [
        ["STR", "Strength", "Melee damage, heavy lifting, how much you can carry."],
        ["CON", "Constitution", "Your Hit Points and resisting cold/disease."],
        ["AGL", "Agility", "Movement, dodging, ranged/finesse damage."],
        ["INT", "Intelligence", "Awareness, lore, healing, and magic skill."],
        ["WIL", "Willpower", "Your Willpower Points and resisting fear."],
        ["CHA", "Charisma", "Persuasion, performance, social skills."],
    ],
    widths=[0.8, 1.6, 4.0],
)

doc.add_heading("The core dice roll: roll UNDER", level=2)
para("Dragonbane runs on one main idea. To attempt something risky you roll a "
     "twenty-sided die (a d20) and try to roll EQUAL TO OR LOWER THAN your "
     "skill level. Low rolls are good.")
bullet("If you roll your skill level or under → success.", bold_lead="")
bullet("If you roll higher → failure.", bold_lead="")
rich([("Roll a natural ", {}), ("1", {"bold": True, "color": GOOD}),
      (" → a ", {}), ("Dragon", {"bold": True, "color": GOOD}),
      (" (a critical success — something extra good happens).", {})])
rich([("Roll a natural ", {}), ("20", {"bold": True, "color": BAD}),
      (" → a ", {}), ("Demon", {"bold": True, "color": BAD}),
      (" (a critical failure — something goes wrong).", {})])
para("The app shows the die face, tells you whether you succeeded, and flags "
     "Dragons and Demons for you.", italic=True, color=MUTED)

doc.add_heading("Boons and banes", level=2)
para("Sometimes a roll is easier or harder. A boon (advantage) or a bane "
     "(disadvantage) means you roll two d20s instead of one:")
bullet("Boon: roll 2 dice, keep the LOWER (better) one.", bold_lead="")
bullet("Bane: roll 2 dice, keep the HIGHER (worse) one.", bold_lead="")
para("Boons and banes can stack, and they cancel each other out one-for-one. "
     "The app tracks the total for you and even applies banes automatically "
     "(for example, when a condition or heavy armour is affecting you).")

doc.add_heading("Pushing a roll", level=2)
para("If you fail a roll, you may 'push' it — try again once — but there's a "
     "price: you take a condition of your choice. The app offers you the push "
     "and lets you pick which condition to suffer.")

doc.add_heading("Conditions", level=2)
para("There are six conditions, one tied to each attribute. While you have a "
     "condition, rolls using its attribute suffer a bane. The app applies this "
     "bane automatically and marks the attribute with a ⚠.")
table(
    ["Condition", "Attribute affected"],
    [
        ["Exhausted", "STR"],
        ["Sickly", "CON"],
        ["Dazed", "AGL"],
        ["Angry", "INT"],
        ["Scared", "WIL"],
        ["Disheartened", "CHA"],
    ],
    widths=[2.2, 2.2],
)

doc.add_heading("Hit Points and Willpower Points", level=2)
bullet("Hit Points (HP) measure your health. They start equal to your CON. Reach 0 and you begin dying.", bold_lead="")
bullet("Willpower Points (WP) fuel your magic and heroic abilities. They start equal to your WIL.", bold_lead="")

callout("You don't have to memorise any of this.",
        "The app calculates your HP, WP, movement, damage bonuses and success "
        "chances the moment you finish creating a hero, and re-checks them "
        "every time something changes.", tint="EAF2E6", edge=GOOD)
pagebreak()

# ==========================================================================
# 4. CREATING YOUR FIRST HERO
# ==========================================================================
doc.add_heading("4.  Creating your first hero", level=1)
para("From the Heroes (⚔) tab, tap ", space_after=2)
rich([("“Forge a new hero.”", {"bold": True}),
      (" This opens the creation wizard — a series of 10 short steps. A "
       "progress line at the top shows 'Step X of 10'. Use ", {}),
      ("Next", {"bold": True}), (" and ", {}), ("Back", {"bold": True}),
      (" to move between steps, or the ", {}),
      ("✕", {"bold": True}),
      (" in the corner to cancel. The wizard won't let you continue a step "
       "until your choices are legal, and it explains what's missing if you "
       "try.", {})])

callout("In a hurry? Use a pre-made hero instead.",
        "If you'd rather skip creation for your first game, tap 'Use a "
        "pre-generated hero' on the Heroes tab (see Section 5). You can always "
        "build your own later.")

doc.add_heading("Step 1 — Attributes", level=2)
para("Tap 'Roll attributes'. The app rolls four six-sided dice and drops the "
     "lowest, six times, giving you six scores. Then, using the drop-down next "
     "to each attribute, assign each rolled score to STR, CON, AGL, INT, WIL "
     "or CHA. Every score must be placed. Not happy with the numbers? Tap "
     "'Re-roll all' to start over.")
bullet("Think about your profession: a Fighter wants a high STR; a Mage wants a high WIL and INT; a Thief wants AGL.", bold_lead="")
bullet("Age modifiers (Step 4) are applied later — the app handles that automatically.", bold_lead="")

doc.add_heading("Step 2 — Kin", level=2)
para("Choose your people. Each kin has a base movement speed and one innate "
     "special ability. Tap a card to select it.")
table(
    ["Kin", "Move", "Innate ability"],
    [
        ["Human", "10", "Adaptive — swap in a different skill for a roll (with the GM's OK)."],
        ["Halfling", "8", "Hard to Catch — a boon when dodging."],
        ["Dwarf", "8", "Unforgiving — a boon attacking someone who has hurt you."],
        ["Elf", "10", "Inner Peace — meditate during a rest to heal extra HP, WP and a condition."],
        ["Mallard", "8", "Ill-Tempered / Webbed Feet — duck-folk grit and swimming."],
        ["Wolfkin", "12", "Hunting Instincts — keen senses and fast movement."],
    ],
    widths=[1.1, 0.7, 4.4],
)

doc.add_heading("Step 3 — Profession", level=2)
para("Your profession is your character class. It sets which skills you're "
     "good at, your starting gear, and a starting heroic ability (except the "
     "Mage, who gets magic instead). Each profession has a 'key attribute' "
     "shown on its card.")
table(
    ["Profession", "Key", "Feel"],
    [
        ["Artisan", "STR", "A skilled crafter."],
        ["Bard", "CHA", "A performer; optionally a Harmonism spellcaster."],
        ["Fighter", "STR", "A frontline warrior."],
        ["Hunter", "AGL", "A ranged and wilderness expert."],
        ["Knight", "STR", "An armoured, honourable combatant."],
        ["Mage", "WIL", "A spellcaster — you also pick a school of magic."],
        ["Mariner", "AGL", "A sailor and adventurer."],
        ["Merchant", "CHA", "A silver-tongued trader."],
        ["Scholar", "INT", "A learned researcher."],
        ["Thief", "AGL", "A stealthy, nimble rogue."],
    ],
    widths=[1.4, 0.6, 4.2],
)
rich([("If you pick ", {}), ("Mage", {"bold": True}),
      (", a second row of cards appears: choose your ", {}),
      ("school of magic", {"bold": True}),
      (" (Animism, Elementalism or Mentalism by default). If you pick ", {}),
      ("Bard", {"bold": True}),
      (" and the Book of Magic content is turned on, you'll see a toggle to "
       "study Harmonism and cast spells using Performance.", {})])

doc.add_heading("Step 4 — Age", level=2)
para("Pick Young, Adult or Old. Age adjusts a few attributes and sets how many "
     "trained skills you get. Older heroes are weaker in body but wiser and "
     "more skilled.")
table(
    ["Age", "Trained skills", "Attribute changes"],
    [
        ["Young", "8", "AGL +1, CON +1"],
        ["Adult", "10", "No change"],
        ["Old", "12", "STR −2, AGL −2, CON −2, INT +1, WIL +1"],
    ],
    widths=[1.2, 1.6, 3.4],
)
para("After you choose an age, the wizard shows a 'Final attributes' box with "
     "your adjusted scores plus your derived HP, WP, Move and damage bonuses.",
     italic=True, color=MUTED)

doc.add_heading("Step 5 — Trained skills", level=2)
para("Skills are things you can attempt, each governed by an attribute. A "
     "TRAINED skill starts at twice its base chance; an untrained skill starts "
     "at just its base chance. Tap skill chips to train them.")
bullet("Pick exactly the number your age allows (8, 10 or 12).", bold_lead="")
bullet("At least 6 must come from your Profession skills (shown in the top group).", bold_lead="")
bullet("The rest are free picks from the Other skills group.", bold_lead="")
para("A live counter at the top tells you how many you've chosen and how many "
     "are from your profession, so you can't get it wrong. (Mages: your magic "
     "school is trained for you automatically.)", italic=True, color=MUTED)

doc.add_heading("Step 6 — Heroic ability (non-mages)", level=2)
para("Warriors and experts choose a heroic ability — a special manoeuvre or "
     "talent, often powered by Willpower Points. Your profession offers one or "
     "two to pick from. Skill requirements are waived for your starting "
     "ability. (Mages skip this step; their power is magic.)")
para("Playing with Solo Mode on? You choose two starting heroic abilities "
     "instead of one.", italic=True, color=MUTED)

doc.add_heading("Step 7 — Magic (casters only)", level=2)
para("If you're a Mage or a Harmonism Bard, this step lets you choose your "
     "opening spells: three magic tricks (minor rank-0 cantrips) and three "
     "rank-1 spells. Mages may pick from their school or from General Magic; "
     "Harmonism Bards pick from Harmonism. Tap chips to select; a counter "
     "keeps you to exactly three of each.")

doc.add_heading("Step 8 — Starting gear", level=2)
para("Choose a starting equipment package, or tap '🎲 Roll a random package' "
     "to let fate decide. Any dice listed in a package (for coins or rations) "
     "are rolled automatically when the hero is created.")

doc.add_heading("Step 9 — Details", level=2)
para("Give your hero a name (required) and, optionally, an appearance, a "
     "weakness, and a memento. Each field has a 🎲 button that suggests a "
     "random, flavourful option if you're stuck. Your weakness and memento "
     "matter later — a weakness can be 'overcome' for bonus advancement, and a "
     "memento is a keepsake that doesn't weigh you down.")

doc.add_heading("Step 10 — Review", level=2)
para("The final step shows a full summary: attributes, derived stats, trained "
     "skills, abilities, magic and gear. If everything looks right, tap "
     "'Create hero'. Your new character is saved and their sheet opens "
     "immediately. If something's off, tap Back to fix it.")

callout("Nothing is permanent.",
        "After creation you can still rename, re-equip, add items, learn "
        "spells, and adjust almost everything from the character sheet. Don't "
        "worry about getting every detail perfect on the first pass.")
pagebreak()

# ==========================================================================
# 5. PRE-GENS
# ==========================================================================
doc.add_heading("5.  Using a ready-made hero", level=1)
para("The fastest way to start playing is to grab one of the five official "
     "Core Set pre-generated characters. On the Heroes tab tap ", space_after=2)
rich([("“Use a pre-generated hero.”", {"bold": True}),
      (" You'll see five cards with a short description and their attributes. "
       "Tap one to add a fully-built copy to your roster and open its sheet.", {})])
table(
    ["Pre-gen", "Kin", "Profession"],
    [
        ["Archmaster Aodhan", "Human", "Mage (Elementalism)"],
        ["Orla Moonsilver", "Elf", "Hunter"],
        ["Makander of Halfbay", "Mallard", "Knight"],
        ["Krisanna the Bold", "Halfling", "Thief"],
        ["Bastonn Bloodjaw", "Wolfkin", "Fighter"],
    ],
    widths=[2.4, 1.2, 2.4],
)
para("A pre-gen is a normal hero once created — rename it, change its gear, or "
     "adjust anything, exactly as you would a hero you built yourself.",
     italic=True, color=MUTED)

# ==========================================================================
# 6. THE CHARACTER SHEET
# ==========================================================================
doc.add_heading("6.  Reading your character sheet", level=1)
para("The character sheet is where you'll spend most of your time in play. "
     "Open it by tapping a hero on the Heroes tab. It scrolls top-to-bottom "
     "through a series of panels. The '← Heroes' button at the top-left takes "
     "you back to your roster. Here is every panel, in order.")

doc.add_heading("Identity & portrait", level=2)
para("The top panel shows your name, kin, profession, school (if any) and age. "
     "Tap the circular portrait to upload a picture from your device — it's "
     "shrunk automatically to keep things fast.")

doc.add_heading("Attributes bar", level=2)
para("Your six attribute scores. If a condition is affecting one, it shows a "
     "⚠ warning — a reminder that rolls using that attribute will take a bane.")
rich([("Just below, a line shows your ", {}), ("Move", {"bold": True}),
      (" speed, your ", {}), ("STR", {"bold": True}), (" and ", {}),
      ("AGL", {"bold": True}), (" damage bonuses, and your ", {}),
      ("encumbrance limit", {"bold": True}),
      (" (how many item slots you can carry).", {})])

doc.add_heading("Hit Points & Willpower (the steppers)", level=2)
para("Two big counters show HP and WP as 'current / maximum'. Tap the − and + "
     "buttons to lower or raise them as you take damage, spend willpower, or "
     "heal. HP and WP can never go below 0 or above your maximum.")
callout("Reaching 0 HP starts the dying rules.",
        "When your HP hits 0, a red 'Dying' panel appears (see below). Pressing "
        "the HP − button again while at 0 counts as a failed death roll — the "
        "app is applying the rule for taking damage while down.", tint="F6E7E2",
        edge=BAD)

doc.add_heading("Movement pool", level=2)
para("A tactical movement tracker for combat. It shows how many metres you "
     "have left this round and offers buttons for common situations: Dash "
     "(double speed, uses your action), Mount, drop Prone / Stand, walk +1/+2/"
     "+4 m, pass through a Door, cross Water, Leap a gap, cross Rough terrain, "
     "and Disengage from an enemy. You can ignore all of this outside combat "
     "— it's here when a fight needs precise positioning.")

doc.add_heading("Rest buttons", level=2)
para("Three quick recovery buttons (explained fully in Section 9):")
bullet("Round rest — regain a little Willpower (once per shift).", bold_lead="")
bullet("Stretch rest — heal some HP and WP and clear one condition (once per shift).", bold_lead="")
bullet("Shift rest — full HP and WP, and clear all conditions.", bold_lead="")

doc.add_heading("Dying panel (only at 0 HP)", level=2)
para("If you're at 0 HP, this panel runs the death-roll sequence for you: "
     "press 'Death roll' each round to roll a d20 against your CON. Three "
     "successes and you stabilise (and recover some HP); three failures and "
     "your hero dies. A Dragon counts as two successes, a Demon as two "
     "failures. You can also 'Rally' to keep acting, or be 'Saved' by an ally's "
     "successful Healing roll.")

doc.add_heading("Conditions", level=2)
para("Six tappable chips for the six conditions. Tap one to toggle it on or "
     "off. While a condition is on, the app automatically banes every roll "
     "that uses its attribute — you don't have to remember to apply it.")

doc.add_heading("Skills", level=2)
para("A list of all your skills with their level and governing attribute. This "
     "panel is the heart of play:")
bullet("Tap a skill's name to roll it (see Section 7). ", bold_lead="")
bullet("The small ◦ / ● dot on the left is an advancement mark — tap it to toggle it. The app ticks it automatically whenever you roll a Dragon or Demon on that skill.", bold_lead="")
bullet("A ⚠ next to a skill means a condition is currently baning it.", bold_lead="")
para("The row of buttons above the list handles getting better between "
     "sessions — 'End session — advancement', 'Train teacher', 'Study "
     "library', 'Gain ability', and more (see Section 9).", italic=True,
     color=MUTED)

doc.add_heading("Abilities", level=2)
para("Lists your kin ability and any heroic abilities, with their Willpower "
     "cost and full rules text so you always know what they do.")

doc.add_heading("Magic & Tricks (casters only)", level=2)
para("Your known magic tricks and spells, each with a 'Cast' button. Casting "
     "is explained in Section 10. There's also a '＋ Learn a spell or school' "
     "button for picking up new magic during play.")

doc.add_heading("Effects, familiar & companions", level=2)
bullet("Active Spells & Effects — track ongoing spells (a '+ Track' button appears on lasting spells you cast).", bold_lead="")
bullet("Familiar (casters) — split some of your Willpower into a familiar with its own pool.", bold_lead="")
bullet("Summons & Companions — track pets, summoned creatures, or raised undead, each with their own HP.", bold_lead="")

doc.add_heading("Inventory & encumbrance", level=2)
para("Dragonbane uses a slot system: you can carry a number of item slots "
     "equal to half your STR (rounded up). The bar at the top of this panel "
     "shows slots used vs your limit; it turns red and prompts a STR roll if "
     "you're over-encumbered.")
bullet("Add items with the 'Add an item…' box (set a weight, then Add).", bold_lead="")
bullet("Equip armour, a helmet, or up to three weapons 'at hand' — equipped gear doesn't count against your slots.", bold_lead="")
bullet("Equipped weapons and shields with a durability rating get a small counter; at 0 they break (💥).", bold_lead="")
bullet("A ⚔ button next to a weapon rolls an attack/damage with it.", bold_lead="")
bullet("Coins add slots too: every 100 coins is one slot. Potions and brews get a '🧪 Use' button.", bold_lead="")

doc.add_heading("Money", level=2)
para("Steppers for gold, silver and copper coins. Adjust with − and +.")

doc.add_heading("Character & notes", level=2)
para("Your appearance and weakness, a place to overcome or change your "
     "weakness, and a free-form Notes / Journal box for anything you like — "
     "session recaps, plot threads, loot lists. At the very bottom sits the "
     "'Delete hero' button.")
callout("Everything saves itself.",
        "There is no save button. Every change you make — a damage tick, a new "
        "item, a note — is written to your device instantly.", tint="EAF2E6",
        edge=GOOD)
pagebreak()

# ==========================================================================
# 7. ROLLING DICE
# ==========================================================================
doc.add_heading("7.  Rolling dice", level=1)
para("This is the single most important thing to learn, and the app makes it "
     "easy.")

doc.add_heading("Making a skill roll", level=2)
numbered("On your character sheet, tap the name of the skill you want to use.")
numbered("A roll window opens showing your skill level and its attribute. If a condition or heavy armour applies, the app has already set a bane and tells you why.")
numbered("If the situation calls for a boon or an extra bane, tap the + or − next to 'Boon / Bane' to adjust the total. (Usually your GM tells you when.)")
numbered("Tap 'Roll d20'. The app shows the die (or dice), highlights the one that counts, and states Success or Failure.")

doc.add_heading("Dragons and Demons", level=2)
rich([("A natural ", {}), ("1", {"bold": True, "color": GOOD}),
      (" is a ", {}), ("Dragon", {"bold": True, "color": GOOD}),
      (" — a critical success. A natural ", {}),
      ("20", {"bold": True, "color": BAD}), (" is a ", {}),
      ("Demon", {"bold": True, "color": BAD}),
      (" — a critical failure. Either one automatically adds an advancement "
       "mark to that skill, so criticals help you improve it later.", {})])

doc.add_heading("Pushing a failed roll", level=2)
para("If you fail (and it wasn't a Demon), the window offers to let you push. "
     "Pick a condition to suffer, and the app re-rolls for you. You can keep "
     "conditions straight — the app won't let you take the same one twice, and "
     "if you already hold all six it explains you can't push further.")

doc.add_heading("Attacking and dealing damage", level=2)
para("Tap the ⚔ button next to an equipped weapon (on the sheet or in "
     "combat). The app rolls your attack using the correct weapon skill — "
     "applying a bane if you don't meet the weapon's STR requirement, offering "
     "a two-handed grip option, and so on — then, on a hit, rolls the damage "
     "dice WITH your STR or AGL damage bonus already added. If you're in a "
     "fight, a target picker lets you apply that damage straight to a monster, "
     "subtracting its armour for you.")

doc.add_heading("Casting a spell", level=2)
para("Tap 'Cast' on a spell. Choose a power level if the spell can be boosted "
     "(each level costs Willpower), then roll your magic skill. The app deducts "
     "the WP (even on a failure — magic is costly), applies Dragon perks, and "
     "on a Demon rolls the full magical mishap table for you. Magic tricks "
     "cost 1 WP and simply succeed. (More in Section 10.)")

callout("The golden rule of the dice engine:",
        "You never have to add up modifiers yourself. Conditions, armour "
        "banes, damage bonuses, ammunition, willpower costs and advancement "
        "marks are all applied automatically. Just tap and read the result.",
        tint="EAF2E6", edge=GOOD)

# ==========================================================================
# 8. COMBAT
# ==========================================================================
doc.add_heading("8.  Fighting: the Combat tab", level=1)
para("When a fight breaks out, switch to the Combat (🛡) tab. It's a shared "
     "battle dashboard that tracks turn order and everyone's health.")

doc.add_heading("Setting up the fight", level=2)
numbered("Add your heroes from the 'Add a hero…' drop-down.")
numbered("Add enemies: pick from the built-in Bestiary monsters, the Rulebook NPCs & Animals list, or type a custom NPC name.")
numbered("Tap 'Draw initiative'. Dragonbane uses initiative cards numbered 1–10; the app deals them and sorts everyone into turn order (1 acts first).")

doc.add_heading("Running the rounds", level=2)
bullet("Each combatant's row shows their initiative number, HP (and WP), and quick attack buttons.", bold_lead="")
bullet("Tap a hero's ⚔ button to roll an attack; tap a monster's 🎲 to roll its attack from its own table — the app rolls the damage and opens a target picker.", bold_lead="")
bullet("Adjust anyone's HP with the inline − / + steppers. A defeated enemy is greyed out and marked.", bold_lead="")
bullet("Use 'Next turn' to step through the order, and 'Next round' (or 'Re-draw') to start a fresh round with new initiative.", bold_lead="")
bullet("'End combat' clears the tracker when the fight is over.", bold_lead="")
para("Hero HP changes in combat sync back to their character sheets, and vice "
     "versa, so nothing gets out of step.", italic=True, color=MUTED)

callout("Who controls the round?",
        "In a solo or local game, anyone can draw initiative and advance turns. "
        "In a synced campaign, the round controls (initiative, next turn, "
        "reset) are reserved for the GM.")

# ==========================================================================
# 9. REST & ADVANCEMENT
# ==========================================================================
doc.add_heading("9.  Rest, healing & getting better", level=1)

doc.add_heading("Resting", level=2)
para("Use the three rest buttons on your sheet:")
table(
    ["Rest", "Time", "Recovers", "Limit"],
    [
        ["Round rest", "~10 seconds", "+D6 Willpower", "Once per shift"],
        ["Stretch rest", "A few minutes", "+D6 HP and +D6 WP, and heal one condition", "Once per shift"],
        ["Shift rest", "~6 hours", "Full HP and WP; all conditions cleared", "Resets the above"],
    ],
    widths=[1.1, 1.3, 3.0, 1.2],
)
para("A stretch rest can heal even more if an ally tends you with a successful "
     "Healing roll, or if you have the Elf's Inner Peace or the Fast Healer "
     "ability — the app offers these options in the rest window. After a shift "
     "rest, casters get a pop-up to prepare their grimoire spells.",
     italic=True, color=MUTED)

doc.add_heading("Advancing your skills", level=2)
para("You improve by using skills and rolling criticals. Each Dragon or Demon "
     "marks a skill for advancement (the ● dot). At the end of a session, tap "
     "'End session — advancement'. The app asks the five advancement "
     "questions — each 'Yes' lets you mark one more skill — then rolls each "
     "marked skill: roll higher than the skill's current level and it goes up "
     "by 1 (max 18). Reaching 18 earns a free heroic ability.")

doc.add_heading("Other ways to improve", level=2)
bullet("Train teacher — spend downtime with an NPC teacher (skill 15+) for one advancement roll on a chosen skill.", bold_lead="")
bullet("Study library — an advancement roll in Beast Lore, Myths & Legends, or Languages.", bold_lead="")
bullet("Gain ability — pick a new heroic ability; ones whose skill requirement you don't meet are locked.", bold_lead="")
bullet("Overcome Weakness — acting against your weakness earns 2 advancement marks and clears the weakness.", bold_lead="")
bullet("Catch up (Death) — a replacement hero after a death gets bonus advancement to catch up to the party.", bold_lead="")

# ==========================================================================
# 10. MAGIC
# ==========================================================================
doc.add_heading("10.  Magic", level=1)
para("If your hero is a Mage or a Harmonism Bard, magic works like this.")
doc.add_heading("Tricks vs spells", level=2)
bullet("Magic tricks (rank 0) are minor cantrips. They cost 1 WP and always succeed.", bold_lead="")
bullet("Spells (rank 1–3) cost Willpower by power level and require a magic-skill roll to work. Willpower is spent even if the roll fails.", bold_lead="")
doc.add_heading("Casting", level=2)
para("On the Magic panel of your sheet, tap 'Cast' on a spell, choose a power "
     "level (higher levels can boost range or damage, at more WP), and roll. "
     "The app handles the WP cost, Dragon bonuses, and — on a Demon — the "
     "magical mishap table automatically. Some spells add buttons to summon a "
     "creature, craft/brew an item into your inventory, or track an ongoing "
     "effect.")
doc.add_heading("Learning more magic", level=2)
para("Use '＋ Learn a spell or school' in the Magic panel to add new spells or "
     "even a whole new school during play (subject to the usual requirements). "
     "Turning on 'Book of Magic' content in Settings unlocks nine extra "
     "schools and many more spells.")

# ==========================================================================
# 11. SOLO
# ==========================================================================
doc.add_heading("11.  Playing solo (no Game Master)", level=1)
para("Dragonbane can be played entirely on your own. Turn on Solo Mode in the "
     "About (⚙) tab (or from the Solo tab's banner) and a 🧭 Solo tab appears "
     "with a GM-less toolkit:")
bullet("Fortune Chart (Oracle) — ask a yes/no question, set how likely it is, and roll for fate's answer.", bold_lead="")
bullet("Inspiration Table — roll three d20s for a random Action + Attribute + Thing prompt when you're stuck.", bold_lead="")
bullet("Narrative Twists — Dragon and Demon twist rollers for dramatic swings out of combat.", bold_lead="")
bullet("Solo NPC & Foe Generator — spin up an NPC or enemy and drop it straight into the Combat tracker.", bold_lead="")
para("Solo Mode also unlocks the solo-only heroic abilities (Army of One, Sole "
     "Survivor) during character creation, a two-ability starting pick, a "
     "'Mission complete (+5 marks)' advancement button, and a 'fail forward' "
     "option on failed rolls.", italic=True, color=MUTED)

# ==========================================================================
# 12. RULES LIBRARY
# ==========================================================================
doc.add_heading("12.  The Rules Library", level=1)
para("The Rules (📖) tab is a complete, searchable reference. Type in the "
     "search box to filter instantly across every category, or open the "
     "accordions to browse:")
bullet("Core Loop & Gameplay Stages, and Wilderness Journeys & Travel.", bold_lead="")
bullet("All six Kin and their abilities; all ten Professions.", bold_lead="")
bullet("Every Skill with its governing attribute; all Heroic Abilities.", bold_lead="")
bullet("All Spells & Tricks by school; Weapons, Armor and Adventuring Gear with full stats.", bold_lead="")
para("It's the fastest way to look something up mid-game without leaving the "
     "app.", italic=True, color=MUTED)

# ==========================================================================
# 13. SETTINGS & SYNC
# ==========================================================================
doc.add_heading("13.  Settings, saving & playing together", level=1)
doc.add_heading("Themes", level=2)
para("The ☾ / ☀ button in the header toggles light and dark themes any time.")

doc.add_heading("Content toggles (About tab)", level=2)
table(
    ["Toggle", "What it adds"],
    [
        ["Book of Magic", "Nine extra magic schools and many more spells in creation and the Rules Library."],
        ["Solo Mode", "The Solo tab and solo-only heroic abilities."],
        ["Advanced / GM Automation", "A GM panel on the sheet: time clock, sleep, light burn-out, cold/disease, fear, concentration."],
        ["GM Screen", "A 🎲 GM dashboard tab (party overview, drop-in monsters, hand out damage/conditions)."],
    ],
    widths=[1.8, 4.0],
)

doc.add_heading("Backing up your heroes", level=2)
bullet("Export heroes (JSON) — download all your characters as a file to keep safe or move to another device.", bold_lead="")
bullet("Import heroes (JSON) — load characters back in from a file (merges with what you have).", bold_lead="")
bullet("Clear all storage — wipes every locally-saved hero. Use with care; it can't be undone.", bold_lead="")
callout("Your heroes live on this device.",
        "Because everything is stored locally, clearing your browser data — or "
        "using a different device — means you won't see them unless you've "
        "Exported a backup or set up cloud sync.", tint="F6E7E2", edge=BAD)

doc.add_heading("Playing together with cloud sync", level=2)
para("To share a live party across several devices, the app can connect to a "
     "Firebase backend (your group sets this up once). In the About tab you "
     "can then:")
bullet("Create a campaign — you become its GM and get a memorable join code (e.g. red-dragon-sword).", bold_lead="")
bullet("Join a campaign — enter a friend's code to join their party.", bold_lead="")
bullet("Link a Google account — optional, to back your characters up across devices.", bold_lead="")
para("Once synced, the Heroes and Combat screens show a live party roster with "
     "everyone's HP, WP and conditions, and the combat tracker updates for all "
     "players in real time. Without cloud sync configured, the app simply runs "
     "in private local mode — which is perfect for a single player or a group "
     "sharing one screen.", italic=True, color=MUTED)
pagebreak()

# ==========================================================================
# 14. QUICK REFERENCE
# ==========================================================================
doc.add_heading("14.  Quick reference", level=1)

doc.add_heading("Derived-stat tables", level=2)
para("These are the numbers the app calculates for you automatically.",
     italic=True, color=MUTED)

para("Skill base chance (from the governing attribute):", bold=True, space_after=2)
table(
    ["Attribute score", "1–5", "6–8", "9–12", "13–15", "16–18"],
    [["Base chance", "3", "4", "5", "6", "7"]],
    widths=[1.8, 0.9, 0.9, 0.9, 0.9, 0.9],
)
para("A trained skill starts at twice the base chance; untrained equals the "
     "base chance.", italic=True, color=MUTED, size=9)

para("Damage bonus (from STR for melee, AGL for ranged/finesse):", bold=True,
     space_after=2)
table(
    ["Attribute score", "1–12", "13–16", "17–18"],
    [["Damage bonus", "none", "+D4", "+D6"]],
    widths=[2.0, 1.4, 1.4, 1.4],
)

para("Movement modifier (from AGL, added to your kin's base move):", bold=True,
     space_after=2)
table(
    ["AGL score", "1–6", "7–9", "10–12", "13–15", "16–18"],
    [["Move modifier", "−4", "−2", "0", "+2", "+4"]],
    widths=[1.6, 0.9, 0.9, 0.9, 0.9, 0.9],
)

para("Vitals:", bold=True, space_after=2)
bullet("Maximum Hit Points = your CON.", bold_lead="")
bullet("Maximum Willpower Points = your WIL.", bold_lead="")
bullet("Encumbrance limit = half your STR, rounded up (in item slots).", bold_lead="")

doc.add_heading("Glossary for newcomers", level=2)
gloss = [
    ("Attribute", "One of your six core scores (STR, CON, AGL, INT, WIL, CHA), 3–18."),
    ("Skill", "A learned ability you roll for; each is governed by an attribute."),
    ("Trained skill", "A skill you're good at; starts at double its base chance."),
    ("d20 / roll-under", "The twenty-sided die; roll equal to or under your skill to succeed."),
    ("Dragon", "A natural 1 — a critical success."),
    ("Demon", "A natural 20 — a critical failure (fumble)."),
    ("Boon / Bane", "Roll two d20s; keep the lower (boon) or higher (bane)."),
    ("Push", "Re-roll a failed roll by taking a condition."),
    ("Condition", "One of six negative states; banes rolls using its attribute."),
    ("HP", "Hit Points — your health; 0 means dying."),
    ("WP", "Willpower Points — fuel for magic and heroic abilities."),
    ("Heroic ability", "A special talent or manoeuvre, often costing WP."),
    ("Kin", "Your character's people (Human, Elf, Dwarf, etc.)."),
    ("Encumbrance", "How much you can carry, measured in item slots."),
    ("Round / Stretch / Shift", "Time scales: ~10 seconds / a few minutes / ~6 hours."),
    ("Advancement mark", "A tick on a skill (from a Dragon/Demon) that lets it improve later."),
]
table(["Term", "Meaning"], gloss, widths=[1.7, 4.4])

doc.add_heading("Ten tips for your first session", level=2)
tips = [
    "Start with a pre-generated hero if you're brand new — you'll learn the sheet faster.",
    "Remember the golden rule: low d20 rolls are good (roll under your skill).",
    "Tap a skill's name to roll it — don't calculate anything by hand.",
    "Watch for the ⚠ marks; they warn you a condition is making a roll harder.",
    "Use the rest buttons between fights — Willpower and HP come back surprisingly fast.",
    "Don't fear failure: pushing a roll turns a miss into a second chance.",
    "Equip your armour, helmet and weapons so they don't eat your carry slots.",
    "Keep a running Journal in the notes box — future-you will thank you.",
    "Export a backup of your heroes now and then, especially before clearing browser data.",
    "Explore the Rules tab whenever you're curious — everything is looked up for you.",
]
for tip in tips:
    numbered(tip)

hrule()
para("Now go forge a hero and step into the Misty Vale. May your rolls be low "
     "and your Dragons many.", italic=True, color=ACCENT,
     align=WD_ALIGN_PARAGRAPH.CENTER)
para("Dragonbane Player — Beginner's Manual", size=8, color=MUTED,
     align=WD_ALIGN_PARAGRAPH.CENTER)

# ---- Save -----------------------------------------------------------------
import os
out = os.path.join(os.path.dirname(__file__), "Dragonbane-Player-Beginners-Manual.docx")
doc.save(out)
print("Wrote", out)
